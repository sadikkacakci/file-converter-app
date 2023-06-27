import img2pdf
import os
from docx2pdf import convert
from pdf2docx import Converter
from PyPDF2 import PdfFileMerger,PdfMerger
from docx import Document

def getFileName(path):
    file_name = os.path.basename(path)
    file_name = file_name.split(".")[0]
    return file_name


def convertImage2Pdf(input_list, output_path):
    with open(output_path, "wb") as pdf_file:
        pdf_file.write(img2pdf.convert(input_list))
    print("Successfully Completed!")


def convertDocx2Pdf(input_path,output_path):
    if len(input_path) == 1:
        convert(input_path[0],output_path)
    else: # birden fazla word seçimi için mergepdf yapılıp ardından word'e çevrilebilir.
        assistantMergeDocxs(input_path)
        convert("merged.docx",output_path)
        os.remove("merged.docx")

    
def convertPdf2Docx(input_path, output_path):
    if len(input_path) == 1:
        converter = Converter(input_path[0])
        converter.convert(output_path, start = 0, end = None)
        converter.close()
    else: # birden fazla pdf seçimi için mergepdf yapılıp ardından word'e çevrilebilir.
        print("length",len(input_path))
        assistantMergePdfs(input_path)
        converter = Converter("merged.pdf")
        converter.convert(output_path, start = 0, end = None)
        converter.close()    
        os.remove("merged.pdf")

def mergePdfs(input_paths, output_path):
    merger = PdfFileMerger()
    for path in input_paths:
        merger.append(path)
    merger.write(output_path)
    merger.close()

def assistantMergePdfs(input_paths):
    merger = PdfMerger()
    for path in input_paths:
        merger.append(path)
    merger.write("merged.pdf")
    merger.close()

def assistantMergeDocxs(input_paths):
    merged_document = Document()
    for path in input_paths:
        doc = Document(path)
        for paragraph in doc.paragraphs:
            merged_document.add_paragraph(paragraph.text)
    
    merged_document.save("merged.docx")

